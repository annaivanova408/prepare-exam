#!/usr/bin/env python3
"""Generate per-topic glossary.md (definitions + quiz + flashcards) and web JSON."""

from __future__ import annotations

import json
import re
import shutil
import subprocess
import tempfile
import zipfile
from html import unescape
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from docx import Document

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data"
TOPICS_JSON = DATA_DIR / "topics.json"
EXTRACTED_JSON = DATA_DIR / "extracted_materials.json"

SUPPORTED = {".docx", ".pptx", ".pdf", ".doc", ".txt", ".md"}
SKIP_NAMES = {
    "glossary.md",
    "glossary.json",
}

STOP_PHRASES = {
    "ответ",
    "ответы",
    "вариант",
    "решение",
    "источник",
    "семинар",
    "тренировочные задания",
    "самостоятельная работа",
    "практика",
    "вопрос",
    "задача",
    "пример",
    "какое",
    "какая",
    "какие",
    "если",
    "к категории",
    "утверждений",
    "не является",
    "верно",
    "выберите",
    "рассчитайте",
    "определите",
    "укажите",
    "niu",
    "ниу",
    "https",
    "storyset",
    "квиз",
    "план курса",
    "система контроля",
    "контрольная работа",
    "работа на семинарах",
    "игра",
    "цена журнала",
    "придумай",
    "ничего не происходит",
    "страны",
    "ввп измеряется",
}

INSTRUCTION_PHRASES = (
    "в каждом задании",
    "прочитайте тест",
    "выберите один правильный",
    "выберите единственный правильный",
    "задание",
    "ниу вшэ",
)

BASE_TERMS = [
    "экономические агенты",
    "альтернативные издержки",
    "кривая производственных возможностей",
    "производственные возможности",
    "рациональный выбор",
    "потребности",
    "блага",
    "свободные блага",
    "экономические блага",
    "ресурсы",
    "экономические ресурсы",
    "собственность",
    "частная собственность",
    "государственная собственность",
    "домохозяйство",
    "фирма",
    "государство",
    "труд",
    "земля",
    "капитал",
    "предпринимательские способности",
    "полезность",
    "общая полезность",
    "предельная полезность",
    "спрос",
    "индивидуальный спрос",
    "рыночный спрос",
    "закон спроса",
    "эластичность",
    "эластичность спроса",
    "перекрестная эластичность спроса",
    "эластичность спроса по доходу",
    "предложение",
    "величина предложения",
    "рыночное равновесие",
    "равновесная цена",
    "конкуренция",
    "совершенная конкуренция",
    "несовершенная конкуренция",
    "монополия",
    "олигополия",
    "монополистическая конкуренция",
    "барьеры входа",
    "издержки",
    "постоянные издержки",
    "переменные издержки",
    "предельные издержки",
    "выручка",
    "прибыль",
    "рынок",
    "внешние эффекты",
    "общественные блага",
    "рыночный провал",
    "субсидия",
    "налоги",
    "налоговая нагрузка",
    "неравенство",
    "коэффициент Джини",
    "кривая Лоренца",
    "бедность",
    "благосостояние",
    "ВВП",
    "реальный ВВП",
    "номинальный ВВП",
    "экономический рост",
    "экономический цикл",
    "безработица",
    "естественный уровень безработицы",
    "циклическая безработица",
    "структурная безработица",
    "фрикционная безработица",
    "инфляция",
    "деньги",
    "денежная масса",
    "функции денег",
    "ставка процента",
    "индекс потребительских цен",
    "государственная политика",
    "фискальная политика",
    "монетарная политика",
    "денежно-кредитная политика",
]

TERM_ALIASES = {
    "альтернативнаястоимость": "Альтернативные издержки",
    "оpportunitycost": "Альтернативные издержки",
    "alternativecost": "Альтернативные издержки",
    "free-rider": "Проблема безбилетника",
    "безбилетник": "Проблема безбилетника",
    "криваяпроизводственныхвозможностей": "Кривая производственных возможностей",
    "эластичность спроса (elasticity of demand)": "Эластичность спроса",
    "коэффициент перекрестной эластичности спроса": "Перекрестная эластичность спроса",
    "коэффициент эластичности спроса по доходу": "Эластичность спроса по доходу",
    "выделяют следующие основные функции денег": "Функции денег",
}

OPTION_KEY_MAP = {"А": "A", "Б": "B", "В": "C", "Г": "D", "Д": "E"}


def normalize_space(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\xa0", " ")).strip()


def clean_line(value: str) -> str:
    value = unescape(value)
    value = re.sub(r"[\u200b\u200e\u200f\u0007]", " ", value)
    value = value.translate(str.maketrans({"—": "-", "–": "-", "−": "-", "→": "->", "Þ": "->"}))
    return cleanup_text(value)


def cleanup_text(value: str) -> str:
    value = normalize_space(value)
    value = value.translate(str.maketrans({"—": "-", "–": "-", "−": "-", "→": "->", "Þ": "->"}))
    value = value.replace("!", ".")
    value = re.sub(r"([.!?])(?=[А-ЯA-ZЁ])", r"\1 ", value)
    value = re.sub(r"([а-яёa-z0-9])(?=[А-ЯЁ][а-яё])", r"\1 ", value)
    value = re.sub(r"-\s*([А-Яа-яA-Za-zЁё])", r"- \1", value)
    value = re.sub(r"\s+([,.;:!?])", r"\1", value)
    value = re.sub(r"([,.;:!?])([^\s\d])", r"\1 \2", value)
    value = re.sub(r":\s*;\s*", ": ", value)
    value = re.sub(r";\s*;", ";", value)
    return normalize_space(value)


def iter_topic_dirs() -> list[Path]:
    return sorted([p for p in ROOT.iterdir() if p.is_dir() and p.name.lower().startswith("тема")], key=natural_key)


def natural_key(path: Path) -> list[Any]:
    return [int(x) if x.isdigit() else x.lower() for x in re.split(r"(\d+)", path.name)]


def topic_id(topic_dir: Path) -> str:
    slug = re.sub(r"[^0-9A-Za-zА-Яа-я]+", "_", topic_dir.name.lower()).strip("_")
    return slug


def list_material_files(topic_dir: Path) -> list[Path]:
    files = []
    for path in topic_dir.iterdir():
        if path.is_file() and path.suffix.lower() in SUPPORTED and path.name.lower() not in SKIP_NAMES:
            if path.name.startswith("~$") or path.name == ".DS_Store":
                continue
            files.append(path)
    return sorted(files, key=natural_key)


def extract_docx(path: Path) -> list[str]:
    doc = Document(str(path))
    lines: list[str] = []
    for p in doc.paragraphs:
        text = clean_line(p.text)
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [clean_line(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                lines.append(" | ".join(cells))
    return lines


def extract_pptx(path: Path) -> list[str]:
    lines: list[str] = []
    ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    with zipfile.ZipFile(path) as zf:
        slide_names = sorted(
            [n for n in zf.namelist() if re.match(r"ppt/slides/slide\d+\.xml$", n)],
            key=lambda n: int(re.search(r"slide(\d+)\.xml", n).group(1)),
        )
        for index, name in enumerate(slide_names, start=1):
            root = ET.fromstring(zf.read(name))
            parts = [clean_line(node.text or "") for node in root.findall(".//a:t", ns)]
            parts = [p for p in parts if p]
            if parts:
                lines.append(f"Слайд {index}: " + " ".join(parts))
        note_names = sorted([n for n in zf.namelist() if n.startswith("ppt/notesSlides/") and n.endswith(".xml")])
        for name in note_names:
            root = ET.fromstring(zf.read(name))
            parts = [clean_line(node.text or "") for node in root.findall(".//a:t", ns)]
            parts = [p for p in parts if p]
            if parts:
                lines.append("Заметки: " + " ".join(parts))
    return lines


def extract_pdf(path: Path) -> list[str]:
    if PdfReader is None:
        return [f"Не удалось извлечь PDF: библиотека pypdf недоступна."]
    lines: list[str] = []
    try:
        reader = PdfReader(str(path))
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            for raw in text.splitlines():
                line = clean_line(raw)
                if line:
                    lines.append(f"Стр. {index}: {line}")
    except Exception as exc:
        lines.append(f"Не удалось извлечь PDF: {exc}")
    return lines


def extract_doc(path: Path) -> list[str]:
    textutil = shutil.which("textutil")
    if not textutil:
        return ["Не удалось извлечь DOC: textutil не найден."]
    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / (path.stem + ".txt")
        try:
            subprocess.run(
                [textutil, "-convert", "txt", "-stdout", str(path)],
                check=True,
                stdout=out.open("w", encoding="utf-8"),
                stderr=subprocess.DEVNULL,
            )
            return [clean_line(line) for line in out.read_text(encoding="utf-8", errors="ignore").splitlines() if clean_line(line)]
        except Exception as exc:
            return [f"Не удалось извлечь DOC: {exc}"]


def extract_text_file(path: Path) -> list[str]:
    return [clean_line(line) for line in path.read_text(encoding="utf-8", errors="ignore").splitlines() if clean_line(line)]


def extract_file(path: Path) -> list[str]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".pptx":
        return extract_pptx(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".doc":
        return extract_doc(path)
    return extract_text_file(path)


def split_sentences(lines: list[str]) -> list[str]:
    text = " ".join(lines)
    parts = re.split(r"(?<=[.!?])\s+(?=[А-ЯA-Z0-9])", text)
    return [normalize_space(p) for p in parts if len(normalize_space(p)) > 20]


def term_is_good(term: str) -> bool:
    term = normalize_space(re.sub(r"^[•\-\d.)\s]+", "", term)).strip(" :;,.")
    low = term.lower()
    words = re.findall(r"[A-Za-zА-Яа-яЁё]+", term)
    if len(term) < 3 or len(term) > 85:
        return False
    if re.match(r"^[A-EАБВГДЕ]\.?\s", term, re.I):
        return False
    if any(x in low for x in STOP_PHRASES):
        return False
    if re.search(r"[=<>∑√∞≈≤≥𝑬𝒅𝛼]|https?://|www\.", term):
        return False
    if re.search(r",\s*руб", low):
        return False
    if term.endswith(",") or term.endswith(";"):
        return False
    if any(len(word) > 24 for word in words):
        return False
    if len(words) == 1 and len(words[0]) > 18:
        return False
    if term.count("(") != term.count(")"):
        return False
    if sum(ch.isalpha() for ch in term) < 3:
        return False
    if term.count(" ") > 7:
        return False
    return True


def normalize_term(term: str) -> str:
    term = normalize_space(re.sub(r"^[•\-\d.)\s]+", "", term)).strip(" :;,.\"«»")
    compact = re.sub(r"\s+", "", term.lower())
    low = term.lower()
    for needle, replacement in TERM_ALIASES.items():
        if needle in low or needle in compact:
            return replacement
    return term[:1].upper() + term[1:]


def allowed_terms_for_topic(topic_dir: Path) -> set[str]:
    title = topic_dir.name.lower()
    terms: set[str] = set()
    if "субъекты" in title:
        terms.update([
            "экономические агенты",
            "альтернативные издержки",
            "кривая производственных возможностей",
            "производственные возможности",
            "рациональный выбор",
            "потребности",
            "блага",
            "свободные блага",
            "экономические блага",
            "ресурсы",
            "экономические ресурсы",
            "собственность",
            "частная собственность",
            "государственная собственность",
            "домохозяйство",
            "фирма",
            "государство",
            "труд",
            "земля",
            "капитал",
            "предпринимательские способности",
            "проблема безбилетника",
            "экономическая система",
        ])
    if "индивид" in title:
        terms.update([
            "полезность",
            "общая полезность",
            "предельная полезность",
            "спрос",
            "индивидуальный спрос",
            "рыночный спрос",
            "закон спроса",
            "эластичность",
            "эластичность спроса",
            "перекрестная эластичность спроса",
            "эластичность спроса по доходу",
            "суверенитет потребителя",
            "реальный доход",
            "номинальный доход",
        ])
    if "фирма" in title:
        terms.update([
            "фирма",
            "предложение",
            "величина предложения",
            "конкуренция",
            "совершенная конкуренция",
            "несовершенная конкуренция",
            "монополия",
            "олигополия",
            "монополистическая конкуренция",
            "барьеры входа",
            "издержки",
            "постоянные издержки",
            "переменные издержки",
            "предельные издержки",
            "выручка",
            "прибыль",
            "добросовестная конкуренция",
            "недобросовестная конкуренция",
            "неценовая конкуренция",
        ])
    if "государство" in title:
        terms.update([
            "рынок",
            "государство",
            "спрос",
            "предложение",
            "рыночное равновесие",
            "равновесная цена",
            "излишек потребителя",
            "излишек производителя",
            "внешние эффекты",
            "общественные блага",
            "рыночный провал",
            "субсидия",
            "налоги",
        ])
    if "неравенство" in title:
        terms.update([
            "неравенство",
            "коэффициент Джини",
            "кривая Лоренца",
            "бедность",
            "благосостояние",
            "доходы",
            "трансферты",
            "прогрессивный налог",
            "налоги",
        ])
    if "рост" in title:
        terms.update([
            "ВВП",
            "реальный ВВП",
            "номинальный ВВП",
            "экономический рост",
            "экономический цикл",
            "безработица",
            "естественный уровень безработицы",
            "циклическая безработица",
            "структурная безработица",
            "фрикционная безработица",
            "рецессия",
            "депрессия",
            "бум",
            "оживление",
        ])
    if "деньги" in title:
        terms.update([
            "инфляция",
            "деньги",
            "денежная масса",
            "функции денег",
            "ставка процента",
            "индекс потребительских цен",
            "номинальный ВВП",
            "простые проценты",
            "сложные проценты",
            "потребительская корзина",
        ])
    if "политик" in title:
        terms.update([
            "государственная политика",
            "фискальная политика",
            "монетарная политика",
            "денежно-кредитная политика",
            "налоги",
            "денежная масса",
        ])
    if not terms:
        terms = set(BASE_TERMS)
    return {t.lower() for t in terms}


def term_allowed(term: str, allowed: set[str]) -> bool:
    low = term.lower()
    if low in allowed:
        return True
    return any(low == t or low in t or t in low for t in allowed)


def clean_definition_text(text: str) -> str:
    text = re.sub(r"^(Слайд|Стр\.)\s*\d+:\s*", "", text)
    text = re.sub(r"^\d+\s*", "", text)
    text = text.replace("•", "; ")
    text = cleanup_text(text)
    text = re.sub(r"([А-ЯЁ][а-яё]+):\s*;\s*", r"\1: ", text)
    return text.strip(" ;")


def definition_quality_bad(definition: str) -> bool:
    low = definition.lower()
    if not definition or definition == "требует дополнения":
        return False
    if any(x in low for x in ["№ темы", "smartlms", "онлайн-курс", "департамент теоретической экономики"]):
        return True
    if re.search(r"тема\s*\d", low) and len(definition) < 90:
        return True
    if definition.strip().startswith(":") or definition.strip().lower().endswith("стр."):
        return True
    words = re.findall(r"[A-Za-zА-Яа-яЁё]+", definition)
    if any(len(word) > 38 for word in words):
        return True
    return False


def polish_definition(term: str, definition: str) -> str:
    low = term.lower()
    if low == "закон спроса" and ("qх" in definition.lower() or "px" in definition.lower() or "þ" in definition.lower()):
        return "при прочих равных условиях рост цены снижает величину спроса, а снижение цены увеличивает величину спроса."
    if low == "ресурсы" and "этофакторы" in definition.lower():
        return "это факторы производства, которые используются для создания благ."
    if low == "блага":
        definition = re.sub(r"Блага:\s*", "Виды: ", definition)
    return cleanup_text(definition)


def definition_candidates(source: str, lines: list[str]) -> list[dict[str, Any]]:
    candidates: list[dict[str, Any]] = []
    patterns = [
        re.compile(r"^(?P<term>[^|]{2,80})\s*\|\s*(?P<definition>[^|]{18,360})(?:\||$)"),
        re.compile(r"^(?P<term>[А-ЯA-ZЁ][^:—–-]{2,80})\s*[:—–-]\s*(?P<definition>.{18,360})$"),
        re.compile(r"^(?P<term>[А-ЯA-ZЁ][^.!?]{2,80}?)\s+[-–—]?\s*это\s+(?P<definition>.{18,360})$", re.I),
        re.compile(r"^(?P<definition>.{18,260}?)\s+называется\s+(?P<term>[А-ЯA-ZЁа-яё][^.!?]{2,80})$", re.I),
    ]
    for line in lines:
        compact = re.sub(r"^(Слайд|Стр\.)\s*\d+:\s*", "", line)
        if re.match(r"^[A-DАБВГ]\.?\s", compact, re.I) or "https://" in compact.lower():
            continue
        for pattern in patterns:
            match = pattern.match(compact)
            if not match:
                continue
            term = normalize_term(match.group("term"))
            definition = clean_definition_text(match.group("definition"))
            if term_is_good(term) and 18 <= len(definition) <= 380 and not definition_quality_bad(definition) and not re.match(r"^[A-DАБВГ]\.?\s", definition, re.I):
                candidates.append(
                    {
                        "term": term,
                        "definition": polish_definition(term, definition),
                        "source": source,
                        "inferred": False,
                    }
                )
    return candidates


def known_phrase_candidates(source: str, lines: list[str], allowed: set[str]) -> list[dict[str, Any]]:
    phrases = sorted(allowed, key=len, reverse=True)
    text = " ".join(lines).lower()
    sentences = split_sentences(lines)
    out = []
    for phrase in phrases:
        if phrase in text:
            definition = "требует дополнения"
            for sentence in sentences:
                if phrase in sentence.lower():
                    definition = polish_definition(phrase, clean_definition_text(sentence[:320]))
                    break
            if definition_quality_bad(definition):
                definition = "требует дополнения"
            out.append(
                {
                    "term": normalize_term(phrase),
                    "definition": definition,
                    "source": source,
                    "inferred": True,
                }
            )
    return out


def build_definitions(topic_dir: Path, extracted: dict[str, list[str]]) -> list[dict[str, str]]:
    """Return definitions found in topic materials.

    We prefer explicit definition patterns, and only then fall back to key phrases
    that are present in the materials (with a short contextual sentence).
    """
    by_key: dict[str, dict[str, str]] = {}
    allowed = allowed_terms_for_topic(topic_dir)
    for source, lines in extracted.items():
        for item in definition_candidates(source, lines):
            if not term_is_good(item["term"]):
                continue
            if item["term"].strip().lower() == "free" and "free-rider" in item["definition"].lower():
                continue
            is_glossary_source = "глоссар" in Path(source).name.lower()
            if not is_glossary_source and not term_allowed(item["term"], allowed):
                continue
            key = item["term"].lower()
            existing = by_key.get(key)
            if existing is None or (is_glossary_source and "глоссар" not in Path(existing["source"]).name.lower()):
                by_key[key] = {"term": item["term"], "definition": item["definition"], "source": source}
    for source, lines in extracted.items():
        for item in known_phrase_candidates(source, lines, allowed):
            if not term_is_good(item["term"]):
                continue
            key = item["term"].lower()
            if key not in by_key:
                by_key[key] = {"term": item["term"], "definition": item["definition"], "source": source}
    definitions = list(by_key.values())
    definitions.sort(key=lambda x: (x["definition"] == "требует дополнения", x["term"].lower()))
    return definitions[:24]


def extract_mcq(extracted: dict[str, list[str]]) -> list[dict[str, Any]]:
    """Extract multiple-choice questions when they are present in materials."""
    out: list[dict[str, Any]] = []
    option_re = re.compile(r"^(?:[A-EА-Д])[\).]\s*(.+)$")
    answer_re = re.compile(r"^(?:правильный ответ|ответ)\s*[:\-]\s*([A-EА-Д])\b", re.I)

    quiz_sources = [(source, lines) for source, lines in extracted.items() if "тренировоч" in Path(source).name.lower()]
    if not quiz_sources:
        quiz_sources = list(extracted.items())

    for source, lines in quiz_sources:
        compact_lines = expand_quiz_lines(lines)
        answer_key = extract_answer_key(compact_lines)
        i = 0
        local_question_index = 0
        while i < len(compact_lines) and len(out) < 12:
            line = compact_lines[i].strip()
            if not line or len(line) < 8:
                i += 1
                continue
            if option_re.match(line):
                i += 1
                continue
            lower_line = line.lower()
            if lower_line.startswith(("задание", "ответ", "тренировочные")):
                i += 1
                continue
            numbered = re.match(r"^(?P<number>\d{1,2})[.)]\s+(?P<text>.+)$", line)
            has_options_after = i + 1 < len(compact_lines) and option_re.match(compact_lines[i + 1].strip())
            if not (line.endswith("?") or line.endswith(":") or numbered or has_options_after):
                i += 1
                continue
            local_question_index += 1
            question_number = int(numbered.group("number")) if numbered else local_question_index
            question = numbered.group("text").strip() if numbered else line.strip()
            if any(phrase in question.lower() for phrase in INSTRUCTION_PHRASES):
                i += 1
                continue
            opts: dict[str, str] = {}
            j = i + 1
            while j < len(compact_lines) and len(opts) < 5:
                raw = compact_lines[j].strip()
                m = option_re.match(raw)
                if m:
                    key = raw[0].upper()
                    key = OPTION_KEY_MAP.get(key, key)
                    opts[key] = normalize_space(m.group(1))
                    j += 1
                    continue
                if raw and raw.endswith("?") and len(opts) >= 2:
                    break
                if raw and len(opts) >= 2 and re.match(r"^\d{1,2}[.)]\s+", raw):
                    break
                if raw and len(opts) >= 2 and answer_re.match(raw):
                    break
                if raw and len(opts) >= 2 and not m:
                    # allow a blank or short line between options blocks, but don't drift too far
                    if len(raw) > 120:
                        break
                j += 1
            if len(opts) >= 3:
                answer = answer_key.get(question_number, "требует дополнения")
                k = j
                while k < min(j + 6, len(compact_lines)):
                    a = answer_re.match(compact_lines[k].strip())
                    if a:
                        answer = a.group(1).upper()
                        answer = OPTION_KEY_MAP.get(answer, answer)
                        break
                    k += 1
                out.append({"question": question, "options": opts, "answer": answer, "source": source})
                i = j
            else:
                i += 1
    return out


def expand_quiz_lines(lines: list[str]) -> list[str]:
    expanded: list[str] = []
    for line in lines:
        text = re.sub(r"^(Слайд|Стр\.)\s*\d+:\s*", "", line)
        text = text.replace("\u0007", "|")
        text = re.sub(r"(?<!\d)(\d{1,2}[.)]\s*)", r"\n\1", text)
        text = re.sub(r"([A-EА-Д][.)]\s*)", r"\n\1", text)
        for part in text.splitlines():
            part = normalize_space(part)
            if part:
                expanded.append(part)
    return expanded


def extract_answer_key(lines: list[str]) -> dict[int, str]:
    answers: dict[int, str] = {}
    answer_block_re = re.compile(r"^ответы(?:\s+на\s+тестовые\s+задания)?[:.]?$", re.I)
    letters_re = re.compile(r"^[A-EА-Д](?:\s*[|,;]\s*[A-EА-Д])+$", re.I)
    numbers_re = re.compile(r"^\d+(?:\s*[|,;]\s*\d+)+$")
    for index in range(len(lines) - 1):
        nums_line = lines[index].strip()
        letters_line = lines[index + 1].strip()
        if numbers_re.match(nums_line) and letters_re.match(letters_line):
            nums = [int(x) for x in re.split(r"\s*[|,;]\s*", nums_line) if x.isdigit()]
            letters = [OPTION_KEY_MAP.get(x.upper(), x.upper()) for x in re.split(r"\s*[|,;]\s*", letters_line) if x]
            answers.update({num: letter for num, letter in zip(nums, letters)})
        if "|" in nums_line and not answers:
            parts = [p for p in re.split(r"\s*\|\s*", nums_line) if p]
            split_at = next((i for i, part in enumerate(parts) if re.match(r"^[A-EА-Д]$", part, re.I)), None)
            if split_at:
                nums = [int(x) for x in parts[:split_at] if x.isdigit()]
                letters = [OPTION_KEY_MAP.get(x.upper(), x.upper()) for x in parts[split_at:] if re.match(r"^[A-EА-Д]$", x, re.I)]
                answers.update({num: letter for num, letter in zip(nums, letters)})
        combined = re.findall(r"\d+|[A-EА-Д]", nums_line, flags=re.I)
        split_at = next((i for i, part in enumerate(combined) if re.match(r"^[A-EА-Д]$", part, re.I)), None)
        if split_at and split_at >= 3:
            nums = [int(x) for x in combined[:split_at] if x.isdigit()]
            letters = [OPTION_KEY_MAP.get(x.upper(), x.upper()) for x in combined[split_at:] if re.match(r"^[A-EА-Д]$", x, re.I)]
            answers.update({num: letter for num, letter in zip(nums, letters)})
    for index, line in enumerate(lines):
        if not answer_block_re.match(line.strip()):
            continue
        for offset in range(1, 6):
            if index + offset + 1 >= len(lines):
                break
            nums_line = lines[index + offset].strip()
            letters_line = lines[index + offset + 1].strip()
            if not numbers_re.match(nums_line) or not letters_re.match(letters_line):
                continue
            nums = [int(x) for x in re.split(r"\s*[|,;]\s*", nums_line) if x.isdigit()]
            letters = [OPTION_KEY_MAP.get(x.upper(), x.upper()) for x in re.split(r"\s*[|,;]\s*", letters_line) if x]
            answers.update({num: letter for num, letter in zip(nums, letters)})
            break
    return answers


def build_flashcards(definitions: list[dict[str, str]]) -> list[dict[str, str]]:
    cards = []
    for d in definitions[:12]:
        if not d["definition"] or d["definition"].lower().startswith("требует"):
            continue
        cards.append({"q": f"Что такое {d['term']}?", "a": d["definition"]})
    return cards[:12]


def clean_quiz(topic_dir: Path, quiz: list[dict[str, Any]]) -> list[dict[str, Any]]:
    title = topic_dir.name.lower()
    cleaned = []
    for item in quiz:
        question = item["question"].strip()
        if len(question) < 12:
            continue
        if question.lower() in title or title in question.lower():
            continue
        if any(phrase in question.lower() for phrase in INSTRUCTION_PHRASES):
            continue
        cleaned.append(item)

    if "государство" in title and cleaned:
        answers = ["D", "B", "D", "C", "B", "A", "B", "C", "B"]
        for item, answer in zip(cleaned, answers):
            if item.get("answer") == "требует дополнения":
                item["answer"] = answer
    return cleaned[:10]


def build_fallback_quiz(definitions: list[dict[str, str]]) -> list[dict[str, Any]]:
    usable = [d for d in definitions if d["definition"] and d["definition"] != "требует дополнения"]
    if not usable or len(definitions) < 4:
        return []
    quiz = []
    for index, item in enumerate(usable[:5]):
        distractors = [d for d in definitions if d["term"] != item["term"]][:3]
        options = {"A": item["term"]}
        for key, distractor in zip(["B", "C", "D"], distractors):
            options[key] = distractor["term"]
        quiz.append(
            {
                "question": f"Какой термин соответствует определению: {item['definition']}",
                "options": options,
                "answer": "A",
                "source": item["source"],
            }
        )
    return quiz


def write_glossary_md(topic_dir: Path, title: str, definitions: list[dict[str, str]], quiz: list[dict[str, Any]], cards: list[dict[str, str]]) -> None:
    md: list[str] = [f"# {title}", "", "## Определения", ""]
    if definitions:
        for item in definitions:
            md.extend(
                [
                    f"### {item['term']}",
                    f"{item['definition'] or 'требует дополнения'}",
                    "",
                    f"Источник: {Path(item['source']).name}",
                    "",
                ]
            )
    else:
        md.append("требует дополнения")
        md.append("")

    md.extend(["## Тест для самопроверки", ""])
    if quiz:
        for idx, q in enumerate(quiz, start=1):
            md.append(f"{idx}. {q['question']}")
            for key in ["A", "B", "C", "D", "E"]:
                if key in q["options"]:
                    md.append(f"   - {key}. {q['options'][key]}")
            md.append("")
            md.append(f"Правильный ответ: {q['answer']}")
            md.append("")
    else:
        md.append("1. требует дополнения")
        md.append("   - A. ...")
        md.append("   - B. ...")
        md.append("   - C. ...")
        md.append("   - D. ...")
        md.append("")
        md.append("Правильный ответ: требует дополнения")
        md.append("")

    md.extend(["## Карточки для повторения", ""])
    if cards:
        for idx, card in enumerate(cards, start=1):
            md.append(f"### Карточка {idx}")
            md.append(f"Вопрос: {card['q']}")
            md.append(f"Ответ: {card['a']}")
            md.append("")
    else:
        md.append("требует дополнения")
        md.append("")

    (topic_dir / "glossary.md").write_text("\n".join(md), encoding="utf-8")


def describe_topic(topic: dict[str, Any]) -> str:
    terms = [item["term"] for item in topic.get("definitions", [])[:6]]
    if terms:
        return f"Тема основана на материалах папки «{topic['title']}» и включает определения: {', '.join(terms)}."
    return f"Тема основана на материалах папки «{topic['title']}». Описание требует дополнения."


def main() -> None:
    DATA_DIR.mkdir(exist_ok=True)
    topics: list[dict[str, Any]] = []
    extracted_all: dict[str, Any] = {}

    for topic_dir in iter_topic_dirs():
        extracted: dict[str, list[str]] = {}
        for file_path in list_material_files(topic_dir):
            rel = file_path.relative_to(ROOT).as_posix()
            lines = extract_file(file_path)
            extracted[rel] = lines

        definitions = build_definitions(topic_dir, extracted)
        quiz = clean_quiz(topic_dir, extract_mcq(extracted))
        if not quiz:
            quiz = build_fallback_quiz(definitions)
        cards = build_flashcards(definitions)
        write_glossary_md(topic_dir, topic_dir.name, definitions, quiz, cards)
        topic = {
            "id": topic_id(topic_dir),
            "title": topic_dir.name,
            "path": topic_dir.relative_to(ROOT).as_posix(),
            "description": "",
            "definitions": definitions,
            "quiz": quiz,
            "flashcards": cards,
            "sources": list(extracted.keys()),
        }
        topic["description"] = describe_topic(topic)
        topics.append(topic)
        extracted_all[topic["id"]] = extracted

    TOPICS_JSON.write_text(json.dumps({"topics": topics}, ensure_ascii=False, indent=2), encoding="utf-8")
    EXTRACTED_JSON.write_text(json.dumps(extracted_all, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Generated {len(topics)} topic glossaries")
    print(f"Generated {TOPICS_JSON.relative_to(ROOT)}")
    print("Generated per-topic glossary.md")


if __name__ == "__main__":
    main()
